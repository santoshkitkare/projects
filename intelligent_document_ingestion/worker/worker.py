import os
import json
from datetime import datetime
import tempfile
from concurrent.futures import ThreadPoolExecutor, as_completed

import boto3
from sqlalchemy import (
    create_engine, Column, String, DateTime, Text, Integer, JSON
)
from sqlalchemy.orm import sessionmaker, declarative_base
from dotenv import load_dotenv

import google.generativeai as genai

# extract libs
import PyPDF2
import pytesseract
from PIL import Image
import docx
import pandas as pd

from llm_prompts import build_metadata_prompt


# ================== Env + Config ==================
load_dotenv()

DATABASE_URL = os.getenv("DATABASE_URL")
S3_BUCKET_NAME = os.getenv("S3_BUCKET_NAME")
SQS_QUEUE_URL = os.getenv("SQS_QUEUE_URL")
AWS_REGION = os.getenv("AWS_REGION")
GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# Thread pool for parallel message processing inside one worker process
MAX_WORKER_THREADS = int(os.getenv("WORKER_THREADS", "5"))

print("[WORKER] Starting with config:")
print("  S3 Bucket:", S3_BUCKET_NAME)
print("  AWS Region:", AWS_REGION)
print("  Database URL:", DATABASE_URL)
print("  SQS Queue URL:", SQS_QUEUE_URL)
print("  Worker threads:", MAX_WORKER_THREADS)

# ---------- DB engine with pooling (safe for concurrency) ----------
engine = create_engine(
    DATABASE_URL,
    pool_pre_ping=True,
    pool_size=int(os.getenv("DB_POOL_SIZE", "10")),
    max_overflow=int(os.getenv("DB_MAX_OVERFLOW", "20")),
)
SessionLocal = sessionmaker(bind=engine)
Base = declarative_base()


class Document(Base):
    __tablename__ = "documents"

    file_id = Column(String, primary_key=True, index=True)
    user_id = Column(String, index=True)
    file_name = Column(String)
    file_type = Column(String)
    file_size = Column(Integer)
    status = Column(String, default="pending")  # pending, processing, completed, failed
    upload_time = Column(DateTime, default=datetime.utcnow)
    completed_time = Column(DateTime, nullable=True)
    s3_key = Column(String)
    error = Column(Text, nullable=True)
    extracted_metadata = Column(JSON, nullable=True)  # flexible for demo


# ---------- AWS clients ----------
print("[WORKER] Setting up AWS clients.")
s3_client = boto3.client("s3", region_name=AWS_REGION)
sqs_client = boto3.client("sqs", region_name=AWS_REGION)

# ---------- Gemini model ----------
print("[WORKER] Setting up Gemini model.")
if GOOGLE_API_KEY:
    genai.configure(api_key=GOOGLE_API_KEY)
    # No more list_models() spam – just bind the model directly
    gemini_model = genai.GenerativeModel("gemini-2.5-flash")
else:
    gemini_model = None


# ================== LLM Helpers ==================
def call_gemini_for_json(prompt: str) -> dict:
    """Call Gemini and parse JSON. On failure, return {}."""
    if not gemini_model:
        return {}

    try:
        resp = gemini_model.generate_content(prompt)
        raw = resp.text.strip()
        # In case model wraps JSON in ``` blocks
        if raw.startswith("```"):
            raw = raw.strip("`")
            if raw.lower().startswith("json"):
                raw = raw[4:].strip()
        return json.loads(raw)
    except Exception as e:
        print("[GEMINI ERROR]", e)
        return {}


SUPPORTED_TYPES = [
    "Question Paper",
    "Research Paper",
    "Invoice",
    "Information Document",
]


def classify_document_type(text: str) -> str:
    """Returns one of SUPPORTED_TYPES or 'To be supported'."""
    snippet = text[:4000] if text else ""
    if not snippet or not gemini_model:
        return "To be supported"

    # TODO (optional): add simple heuristics here to avoid LLM call when obvious

    prompt = f"""
You are a strict document classifier.

You must choose EXACTLY ONE of these categories that best describes the document:

1. Question Paper
2. Research Paper
3. Invoice
4. Information Document

If it does not clearly fit any of the above, return exactly:
To be supported

Rules:
- Return ONLY the category string, nothing else.
- Do NOT explain your choice.

Document text (partial):
\"\"\"{snippet}\"\"\""""

    try:
        resp = gemini_model.generate_content(prompt)
        label = resp.text.strip()
    except Exception as e:
        print("[GEMINI CLASSIFICATION ERROR]", e)
        return "To be supported"

    label = label.replace('"', "").strip()
    if label in SUPPORTED_TYPES:
        return label
    return "To be supported"


def extract_structured_metadata(text: str, page_count: int | None) -> dict:
    doc_type = classify_document_type(text)
    if doc_type == "To be supported":
        return {
            "documentType": "To be supported",
            "shortSummary": "Document type not yet supported.",
            "rawTextPreview": text[:500],
            "pageCount": page_count,
        }

    prompt = build_metadata_prompt(doc_type, text, page_count)
    meta = call_gemini_for_json(prompt)
    if not isinstance(meta, dict):
        meta = {}

    # Ensure documentType present
    if "documentType" not in meta:
        meta["documentType"] = doc_type
    meta["pageCount"] = page_count
    return meta


# ================== Extraction Helpers ==================
def extract_pdf_text(fpath):
    with open(fpath, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return {
            "text": text[:5000],   # limit size for DB
            "pageCount": len(reader.pages),
        }


def extract_image_text(fpath):
    img = Image.open(fpath)
    text = pytesseract.image_to_string(img)
    return {
        "text": text[:5000],
        "pageCount": 1,
    }


def extract_docx_text(fpath):
    document = docx.Document(fpath)
    full_text = "\n".join(p.text for p in document.paragraphs)
    return {
        "text": full_text[:5000],
        "pageCount": None,
    }


def extract_csv_info(fpath):
    df = pd.read_csv(fpath, nrows=5)
    return {
        "columns": df.columns.tolist(),
        "sampleRows": df.head(3).to_dict(),
        "pageCount": None,
    }


def run_extraction(file_path: str, file_type: str) -> dict:
    if file_type == "application/pdf":
        return extract_pdf_text(file_path)
    elif file_type in ["image/jpeg", "image/png"]:
        return extract_image_text(file_path)
    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        return extract_docx_text(file_path)
    elif file_type == "text/csv":
        return extract_csv_info(file_path)
    else:
        return {"text": "Unsupported format", "pageCount": None}


# ================== Core Processing ==================
def process_message(body: dict):
    """Process one SQS message: download -> extract -> LLM -> update DB."""
    print("[WORKER] Processing message:", body)
    file_id = body["fileId"]
    s3_info = body["s3Location"]
    bucket = s3_info["bucket"]
    key = s3_info["key"]

    db = SessionLocal()
    doc = db.query(Document).filter(Document.file_id == file_id).first()
    if not doc:
        db.close()
        print("[WORKER] Document not found in DB for file_id:", file_id)
        return

    doc.status = "processing"
    db.commit()

    try:
        # download to temp file
        tmp = tempfile.NamedTemporaryFile(delete=False)
        tmp_path = tmp.name
        tmp.close()

        try:
            s3_client.download_file(Bucket=bucket, Key=key, Filename=tmp_path)
            extracted = run_extraction(tmp_path, doc.file_type)  # {"text": ..., "pageCount": ...}
            text = extracted.get("text", "") or ""
            page_count = extracted.get("pageCount")

            ai_meta = extract_structured_metadata(text, page_count)

            metadata = {
                "processedAt": datetime.utcnow().isoformat() + "Z",
                "documentType": ai_meta.get("documentType", "unknown"),
                "llmMetadata": ai_meta,
                "textPreview": text[:1000],
                "pageCount": page_count,
            }
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)

        doc.status = "completed"
        doc.completed_time = datetime.utcnow()
        doc.extracted_metadata = metadata
        doc.error = None
        db.commit()
        print(f"[OK] Real processing done for {file_id}")

    except Exception as e:
        print(f"[ERROR] {file_id} failed: {e}")
        doc.status = "failed"
        doc.error = str(e)
        db.commit()

    finally:
        db.close()


# ================== SQS Main Loop (with thread pool) ==================
def _handle_sqs_message(msg: dict):
    """Wrapper to process an SQS message (for thread pool)."""
    body = json.loads(msg["Body"])
    process_message(body)
    return msg["ReceiptHandle"]


def main():
    print("[WORKER] Running with REAL extraction + thread pool.")
    print("[WORKER] Listening to SQS queue:", SQS_QUEUE_URL)

    executor = ThreadPoolExecutor(max_workers=MAX_WORKER_THREADS)

    while True:
        resp = sqs_client.receive_message(
            QueueUrl=SQS_QUEUE_URL,
            MaxNumberOfMessages=10,   # pull more per batch
            WaitTimeSeconds=10,       # long-polling to reduce empty calls
        )

        messages = resp.get("Messages", [])
        if not messages:
            continue

        futures = {executor.submit(_handle_sqs_message, m): m for m in messages}

        for fut in as_completed(futures):
            msg = futures[fut]
            try:
                receipt_handle = fut.result()
                # delete only if processing finished without raising
                sqs_client.delete_message(
                    QueueUrl=SQS_QUEUE_URL,
                    ReceiptHandle=receipt_handle,
                )
            except Exception as e:
                # process_message already handles its own errors, so this is defensive
                print("[WORKER] Unexpected error in thread:", e)
                # we do NOT delete the message → SQS will retry after visibility timeout


if __name__ == "__main__":
    print("[WORKER] Starting main loop.")
    main()
